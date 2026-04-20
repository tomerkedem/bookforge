"""
Splits raw paragraphs into chapters based on headings.
Extracts images from Word documents.
Returns a list of chapters with title, content, and images.
"""

import os
import re
import json
from pathlib import Path
from collections import Counter


INTRO_KEYWORDS = ["מבוא", "פתיחה", "הקדמה", "introduction", "preface", "foreword"]
COVER_KEYWORDS = ["שער", "cover", "title"]


# =============================================================================
# Instrumentation: count how many times each regex-cleanup rule actually fires.
# The goal is to confirm which rules are still needed after the ingest.py fix
# that produces cleaner markdown at the source. Rules that never fire across
# several real books are dead code and can be removed.
#
# Usage from build.py:
#     from pipeline.parse import reset_clean_stats, dump_clean_stats
#     reset_clean_stats()
#     ... run pipeline ...
#     dump_clean_stats("clean_markdown_stats.json")
# =============================================================================

_CLEAN_STATS: Counter = Counter()


def reset_clean_stats() -> None:
    """Call at the start of a pipeline run to zero the counters."""
    _CLEAN_STATS.clear()


def dump_clean_stats(path: str = "clean_markdown_stats.json") -> None:
    """Write the accumulated counters to a JSON report."""
    report = {
        "total_substitutions": sum(_CLEAN_STATS.values()),
        "rules_fired": len([v for v in _CLEAN_STATS.values() if v > 0]),
        "by_rule": dict(sorted(_CLEAN_STATS.items(), key=lambda kv: -kv[1])),
    }
    Path(path).write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _counted_sub(rule_id: str, pattern: str, repl, text: str, flags: int = 0) -> str:
    """
    Drop-in replacement for re.sub that also counts how many
    substitutions were made under the given rule_id.
    """
    new_text, count = re.subn(pattern, repl, text, flags=flags)
    if count:
        _CLEAN_STATS[rule_id] += count
    return new_text


def _clean_title(text: str) -> str:
    """
    Clean title/subtitle extracted from cover page.
    Removes HTML tags, markdown formatting, and extra whitespace.
    """
    # Remove HTML tags (e.g., <div style="...">, </div>)
    cleaned = re.sub(r'<[^>]+>', '', text)
    # Remove markdown bold/italic
    cleaned = re.sub(r'\*+', '', cleaned)
    # Remove tabs and normalize whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def extract_book_info(ingested: dict) -> dict:
    """
    Return empty book title and subtitle.

    Book-level metadata (title, subtitle) should be provided via the
    --title CLI argument, not auto-detected from the Word document.
    Auto-detection was unreliable because:
    - Some books use a Word SDT (cover page) element which python-docx
      does not expose as a regular paragraph
    - Some books have no cover page at all and start directly with
      a Heading 1 that is a chapter title, not the book title
    - The previous heuristic (collect paragraphs before first Heading 1)
      grabbed arbitrary content from inside the first chapter when no
      cover existed

    This function is kept as a stub so that call sites in build.py
    and elsewhere continue to work unchanged. The returned values
    are empty, and build.py already falls back to the --title CLI
    argument and to the book slug when title/subtitle are empty.
    """
    return {"title": "", "subtitle": ""}


def _clean_heading(text: str) -> str:
    """
    Remove bold/italic markdown from headings - they're already styled by structure.
    Handles Word artifacts like **פרק 1****:** or ***text***.
    """
    # Remove all asterisk formatting from headings
    # They shouldn't have markdown bold/italic since heading style is already prominent
    cleaned = re.sub(r'\*+', '', text)
    # Clean up extra spaces that may result
    cleaned = re.sub(r'\s{2,}', ' ', cleaned)
    return cleaned.strip()


def _fix_code_blocks(text: str) -> str:
    """
    Fix code block format from Word documents.

    Word format: python``` ... ```
    Target format: ```python ... ```

    Supported languages: python, javascript, bash, markdown, etc.
    Normalizes language names to lowercase.
    Normalizes md -> markdown.

    Instrumented: firings of each sub-rule are counted so we can see
    whether this function is still needed after the ingest.py fix.
    Function is currently not called from anywhere in the codebase,
    so the counts will likely stay at zero.
    """
    # Pattern: language``` at start of line, then code, then ``` at end
    pattern = r'^(\w+)```\s*\n([\s\S]*?)```\s*$'

    def replace_block(match):
        lang = match.group(1).lower()
        if lang == 'md':
            lang = 'markdown'
        code = match.group(2).rstrip()
        _CLEAN_STATS["fix_code_blocks_multiline"] += 1
        return f'```{lang}\n{code}\n```'

    result = re.sub(pattern, replace_block, text, flags=re.MULTILINE)

    # Inline pattern: python```code```
    inline_pattern = r'(\w+)```([^`]+)```'

    def replace_inline(match):
        lang = match.group(1).lower()
        if lang == 'md':
            lang = 'markdown'
        code = match.group(2)
        _CLEAN_STATS["fix_code_blocks_inline"] += 1
        return f'```{lang}\n{code}\n```'

    result = re.sub(inline_pattern, replace_inline, result)

    # Normalize existing language tags (case normalization)
    result = _counted_sub("fix_code_lang_Python", r'```Python\b', '```python', result)
    result = _counted_sub("fix_code_lang_Bash", r'```Bash\b', '```bash', result)
    result = _counted_sub("fix_code_lang_Markdown", r'```Markdown\b', '```markdown', result)
    result = _counted_sub(
        "fix_code_lang_MD_case_insensitive",
        r'```MD\b',
        '```markdown',
        result,
        flags=re.IGNORECASE,
    )
    result = _counted_sub("fix_code_lang_md_lower", r'```md\b', '```markdown', result)

    return result



def _clean_markdown_final(text: str) -> str:
    """
    Instrumented cleanup of markdown artifacts from Word formatting.

    Behavior is identical to the original - same rules in the same order.
    The only change is that every substitution is counted in _CLEAN_STATS
    under a stable rule_id. This lets us see which rules still fire after
    the ingest.py fix, and remove the ones that never do.

    Original patterns (kept for context):
    - **** (4+ asterisks) -> **
    - **:** or **: ** -> :
    - **, ** between bold markers -> ,
    - ** ** (spaced asterisks) -> space
    - Hebrew prefix letter absorbed into bold
    - **text**.** -> **text**.
    - Unbalanced ** -> balanced
    - "- •" or "- 1)" duplicate bullets -> single prefix
    - Corrupted numbering with percent sign -> digit
    - "1) text" at line start -> "**(1)** text"
    - ": **" (colon with space before close) -> ":** "
    """
    # Fix corrupted numbering: %1), %2), %1., %2. etc. -> 1), 2), 1., 2. etc.
    text = _counted_sub("01_percent_digit_marker", r'%(\d+)([\.\)])', r'\1\2', text)

    # Fix duplicate list prefixes: "- •", "- 1)", "- 2." etc.
    text = _counted_sub(
        "02_duplicate_bullet_dash_dot",
        r'^(\s*)-\s+(•)',
        r'\1\2',
        text,
        flags=re.MULTILINE,
    )
    text = _counted_sub(
        "03_duplicate_bullet_dash_number",
        r'^(\s*)-\s+(\d+[\)\.])(\s)',
        r'\1\2\3',
        text,
        flags=re.MULTILINE,
    )

    # Convert standalone "1)" at start of line to "**(1)**"
    text = _counted_sub(
        "04_number_paren_to_bold",
        r'^(\d+)\)\s+',
        r'**(\1)** ',
        text,
        flags=re.MULTILINE,
    )

    # Fix space before closing **
    text = _counted_sub(
        "05_space_before_close_eol",
        r' \*\*$',
        '**',
        text,
        flags=re.MULTILINE,
    )
    text = _counted_sub(
        "06_space_before_close_punct",
        r' \*\*([.,;:!?)])',
        r'**\1',
        text,
    )

    # ": **" -> ":** "
    text = _counted_sub("07_colon_space_before_close", r': \*\*', ':** ', text)

    # Per-line processing
    lines = text.split('\n')
    cleaned_lines = []

    for line in lines:
        # 1. Fix consecutive asterisks (4 or more) -> **
        line = _counted_sub("08_four_plus_asterisks", r'\*{4,}', '**', line)

        # 2. Fix **:** pattern
        line = _counted_sub("09_bold_colon_bold", r'\*\*:\*\*', ':', line)
        line = _counted_sub("10_bold_colon_space_bold", r'\*\*:\s*\*\*', ': ', line)

        # 3. Fix trailing **.**
        line = _counted_sub("11_bold_dot_bold", r'\*\*\.\*\*', '.', line)
        line = _counted_sub("12_bold_dot_eol", r'\*\*\.$', '.', line)

        # 4. Fix **, **
        line = _counted_sub("13_bold_comma_bold", r'\*\*,\s*\*\*', ', ', line)

        # 5. Fix ** ** (space between markers)
        line = _counted_sub("14_bold_space_bold", r'\*\*\s+\*\*', ' ', line)

        # 6. Hebrew prefix before bold
        line = _counted_sub(
            "15_hebrew_prefix_before_bold",
            r'([הובכלמש])\*\*([^*]+)\*\*',
            r'**\1\2**',
            line,
        )

        # 7. Sentence ending with ". **label" without closing
        match = re.search(r'\.\s+\*\*([^*]+)$', line)
        if match and line.count('**') % 2 != 0:
            line = line + '**'
            _CLEAN_STATS["16_close_unbalanced_after_dot"] += 1

        # 8. Bullet with misplaced **
        line = _counted_sub(
            "17_bullet_colon_bold_label",
            r'^(- )([^*]+):\*\*\s*([^*]+)\*\*$',
            r'\1**\2: \3**',
            line,
        )

        # 9. Partial variant of 8
        line = _counted_sub(
            "18_bullet_colon_bold_label_partial",
            r'^(- )([^*:]+):\*\*\s*([^*]+)\*\*',
            r'\1**\2: \3**',
            line,
        )

        # 10. Label at line start
        line = _counted_sub(
            "19_label_colon_at_start",
            r'^(\s*)([^\s*][^*\n]{2,})\*\*:',
            r'\1**\2:**',
            line,
        )

        # 11. Standalone ** on a line
        if re.match(r'^\s*\*\*\s*$', line):
            line = ''
            _CLEAN_STATS["20_standalone_bold_line"] += 1

        # 12. "**text**" inside quotes -> "text"
        line = _counted_sub(
            "21_bold_inside_quotes",
            r'"\*\*([^*"]+)\*\*"',
            r'"\1"',
            line,
        )

        # 13. Remaining unbalanced **
        final_count = line.count('**')
        if final_count % 2 != 0 and final_count > 0:
            if re.search(r'\*\*[^*]+$', line) and not re.search(r'[^*]\*\*$', line):
                line = line + '**'
                _CLEAN_STATS["22_close_trailing_unbalanced"] += 1
            elif re.search(r'^\s*\*\*\s*$', line):
                line = ''
                _CLEAN_STATS["23_remove_lonely_bold"] += 1

        # 14. Collapse multiple spaces
        line = _counted_sub("24_multiple_spaces", r' {2,}', ' ', line)

        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)


def _build_markdown_table(rows: list) -> str:
    """
    Build a markdown pipe-table from a list of rows.

    Each row is a list of cell dicts with a `text` field that already
    contains the cell's content rendered as markdown (with inline
    formatting from ingest.py::_runs_to_markdown). The first row is
    treated as the header; a separator line is emitted after it as
    markdown pipe-tables require.

    Edge cases handled:
    - Empty cells are rendered as a single space so the pipe syntax
      stays valid ("| |" is valid, "||" is not).
    - Rows with differing cell counts are padded with empty cells to
      match the widest row, so the table stays rectangular.
    - Newlines inside a cell are replaced with a space. Markdown pipe
      tables do not support multi-line cells; using <br> would work on
      some renderers but break on others, so a space is the safer
      choice. This is a known limitation of the pipe-table format.
    - Pipe characters inside cell text are already escaped by
      ingest.py (see _extract_table_data), so we don't escape again.
    """
    if not rows:
        return ""

    def cell_text(cell):
        # ingest.py now builds the cell's markdown directly, including
        # inline formatting (bold, italic, code, links). Previously we
        # re-rendered from `runs` here, which lost formatting because
        # the cell-level run extractor used to emit plain text only.
        t = cell.get("text", "")
        # Collapse newlines; pipe tables cannot contain real line breaks
        t = t.replace("\n", " ").replace("\r", " ").strip()
        return t if t else " "

    # Find the widest row so we can pad narrower rows
    max_cols = max(len(row) for row in rows)

    normalized_rows = []
    for row in rows:
        cells = [cell_text(c) for c in row]
        # Pad with empty cells if this row is shorter than the widest
        while len(cells) < max_cols:
            cells.append(" ")
        normalized_rows.append(cells)

    lines = []
    # Header row (first row is assumed to be the header)
    header = normalized_rows[0]
    lines.append("| " + " | ".join(header) + " |")
    # Separator
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    # Body rows
    for cells in normalized_rows[1:]:
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def parse(ingested: dict) -> list[dict]:
    blocks = ingested["blocks"]
    chapters = []
    current = None

    for idx, block in enumerate(blocks):
        block_type = block.get("type")
        style = block.get("style", "")
        text = block.get("text", "")
        doc_idx = block.get("doc_index", idx)

        # Tables: build a proper markdown pipe-table from the structured
        # rows that ingest.py extracted. Previously we passed through an
        # empty `text` field (table blocks have no `text`, only `rows`),
        # which silently dropped every table from the output.
        if block_type == "table":
            if current:
                rows = block.get("rows", [])
                table_md = _build_markdown_table(rows)
                current["content"].append({
                    "text": table_md,
                    "style": "Table",
                    "para_index": doc_idx,
                    "indent_level": 0
                })
            continue

        # We only process paragraph/code/heading-like blocks here
        if block_type not in {"paragraph", "heading", "code"}:
            continue

        if "Heading 1" in style:
            if current:
                chapters.append(current)

            clean_title = _clean_heading(text)
            chapter_type = _classify_chapter(clean_title, len(chapters))
            current = {
                "number": len(chapters) + 1,
                "title": clean_title,
                "heading_doc_index": doc_idx,
                "content": [],
                "has_images": False,
                "type": chapter_type
            }
            continue

        if current:
            if "Heading" in style:
                text = _clean_heading(text)

            current["content"].append({
                "text": text,
                "style": style if block_type != "code" else "code",
                "para_index": doc_idx,
                "indent_level": 0
            })

    if current:
        chapters.append(current)

    return chapters


def _classify_chapter(title: str, index: int) -> str:
    title_lower = title.lower()

    if index == 0 and len(title_lower) < 50:
        for keyword in COVER_KEYWORDS:
            if keyword in title_lower:
                return "cover"

    for keyword in INTRO_KEYWORDS:
        if keyword in title_lower:
            return "intro"

    return "content"


# Default: save images to public/ for web serving (not output/)
# Resolved to absolute path relative to project root
from pathlib import Path as _Path
_PROJECT_ROOT = _Path(__file__).resolve().parent.parent.parent
DEFAULT_ASSETS_DIR = str(_PROJECT_ROOT / "public")


def extract_images(docx_path: str, book_name: str, assets_base_dir: str = DEFAULT_ASSETS_DIR) -> dict:
    """
    Extracts images from a Word document and maps them to paragraph positions.
    
    Images are saved directly to public/{book_name}/assets/ for web serving.
    No duplicate copies in output/ folder.
    
    Cover image detection logic:
    - If an image exists BEFORE the first Heading 1 → that's the cover
    - Otherwise, no cover.png is created (all images numbered sequentially)
    
    Returns a dict with:
      - 'files': mapping rel_id to file path
      - 'positions': list of (paragraph_index, rel_id, filename, width_px, height_px)
      - 'has_cover': boolean indicating if cover was found
      - 'book_name': slug for URL paths
    """
    try:
        from docx import Document
        from docx.oxml import parse_xml
    except ImportError:
        raise ImportError("pip install python-docx")

    import re

    doc = Document(docx_path)
    assets_dir = Path(assets_base_dir) / book_name / "assets"
    assets_dir.mkdir(parents=True, exist_ok=True)

    # Step 1: Find first TWO Heading 1 positions
    # (to handle cases like "פתיחה" + cover image + "Chapter 1")
    first_heading_idx = None
    second_heading_idx = None
    heading_count = 0
    for idx, para in enumerate(doc.paragraphs):
        if para.style and "Heading 1" in para.style.name:
            heading_count += 1
            if heading_count == 1:
                first_heading_idx = idx
            elif heading_count == 2:
                second_heading_idx = idx
                break
    
    # Step 2: Build rel_id to image data mapping
    image_data = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue
        try:
            img_data = rel.target_part.blob
            content_type = rel.target_part.content_type
            ext = content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            image_data[rel_id] = {"data": img_data, "ext": ext}
        except Exception:
            continue

       # Step 3: Check for images in SDT elements (before paragraphs - cover page)
    # Keep tuple shape consistent: (para_idx, run_idx, rel_id, w_px, h_px)
    image_positions_temp = []

    body = doc.element.body
    for element in body:
        tag_name = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag_name == 'sdt':
            blips = element.xpath('.//a:blip', namespaces={
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
            })
            for blip in blips:
                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_id and embed_id in image_data:
                    image_positions_temp.append((-1, 0, embed_id, 0, 0))

    # Step 4: Map images in paragraphs (with dimensions)
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            run_xml = run._element.xml.decode('utf-8') if isinstance(run._element.xml, bytes) else run._element.xml

            if '<w:drawing' in run_xml or '<w:pict' in run_xml or 'r:embed=' in run_xml:
                embeds = re.findall(r'r:embed="(rId\d+)"', run_xml)
                matched_rel_id = None

                for eid in embeds:
                    if eid in image_data:
                        matched_rel_id = eid
                        break

                if not matched_rel_id:
                    for rel_id in image_data.keys():
                        if rel_id in run_xml:
                            matched_rel_id = rel_id
                            break

                if matched_rel_id:
                    w_px, h_px = 0, 0
                    extents = re.findall(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', run_xml)
                    if extents:
                        cx, cy = int(extents[0][0]), int(extents[0][1])
                        w_px = round(cx / 914400 * 96)
                        h_px = round(cy / 914400 * 96)

                    image_positions_temp.append((para_idx, run_idx, matched_rel_id, w_px, h_px))

    # Step 5: Sort consistently by paragraph + run
    image_positions_temp.sort(key=lambda x: (x[0], x[1]))

    # Step 6: Determine cover image
    has_cover = False
    cover_rel_id = None

    if image_positions_temp:
        first_img_para_idx = image_positions_temp[0][0]

        # SDT cover page image
        if first_img_para_idx == -1:
            sdt_images = [img for img in image_positions_temp if img[0] == -1]
            cover_rel_id = sdt_images[-1][2]
            has_cover = True
            print(f"[OK] Cover image found: SDT image (cover page)")

        # Otherwise, first image early in the document is treated as cover
        elif first_img_para_idx < 15:
            cover_rel_id = image_positions_temp[0][2]
            has_cover = True
            print(f"[OK] Cover image found: first image at paragraph {first_img_para_idx}")

        else:
            print(f"[WARN] No cover: First image at para {first_img_para_idx} (too late in document)")
            print("  All images will be numbered sequentially (no cover.png)")

    # DEBUG
    print(f"  [DEBUG] rel images found: {len(image_data)}")
    print(f"  [DEBUG] positioned images found: {len(image_positions_temp)}")
    print(f"  [DEBUG] has_cover: {has_cover}, cover_rel_id: {cover_rel_id}")

    # Step 7: Save images
    image_files = {}
    image_positions = []
    image_counter = 1

    for para_idx, run_idx, rel_id, w_px, h_px in image_positions_temp:
        img_info = image_data[rel_id]

        if has_cover and rel_id == cover_rel_id:
            filename = "cover.png"
            img_path = assets_dir / filename
        else:
            filename = f"image-{str(image_counter).zfill(2)}.{img_info['ext']}"
            img_path = assets_dir / filename
            image_counter += 1

        with open(img_path, "wb") as f:
            f.write(img_info["data"])

        image_files[rel_id] = str(img_path)
        image_positions.append((para_idx, run_idx, rel_id, filename, w_px, h_px))

    print(f"  [DEBUG] saved images: {[p[3] for p in image_positions]}")

    return {
        'files': image_files,
        'positions': image_positions,
        'has_cover': has_cover,
        'book_name': book_name
    }


def _clean_markdown_preserving_code(text: str) -> str:
    """
    Apply markdown cleanup only outside protected code regions.

    Protected regions:
    1. Standard fenced code blocks: ```...```
    2. Full single-line CodeRenderer component lines
    """
    lines = text.splitlines(keepends=True)
    cleaned_lines = []

    in_fenced_code = False

    for line in lines:
        stripped = line.strip()

        # Preserve fenced code blocks exactly as-is
        if stripped.startswith("```"):
            in_fenced_code = not in_fenced_code
            cleaned_lines.append(line)
            continue

        if in_fenced_code:
            cleaned_lines.append(line)
            continue

        # Preserve full CodeRenderer line exactly as-is
        if stripped.startswith("<CodeRenderer ") and stripped.endswith("/>"):
            cleaned_lines.append(line)
            continue

        cleaned_lines.append(_clean_markdown_final(line))

    return "".join(cleaned_lines)

def _preserve_soft_breaks(text: str) -> str:
    """
    Convert Word soft line breaks to Markdown hard line breaks.

    Word Shift+Enter is represented earlier as '\\n'.
    In Markdown, two spaces + newline preserve a visible line break.
    """
    if "\n" not in text:
        return text
    return text.replace("\n", "  \n")

def _detect_code_block_start(text: str) -> str | None:
    """
    Detect standard fenced code block start only.

    Supported:
    - ```python
    - ```bash
    - ```javascript
    - ```

    Instrumented: we count detections to confirm whether this branch
    is still needed after ingest.py classifies code paragraphs by style.
    """
    if not text:
        return None

    stripped = text.strip()
    if not stripped.startswith("```"):
        return None

    lang = stripped[3:].strip().lower()

    aliases = {
        "py": "python",
        "python": "python",
        "sh": "bash",
        "shell": "bash",
        "bash": "bash",
        "zsh": "bash",
        "js": "javascript",
        "javascript": "javascript",
        "ts": "typescript",
        "typescript": "typescript",
        "md": "markdown",
        "markdown": "markdown",
        "yml": "yaml",
        "yaml": "yaml",
        "ps1": "powershell",
        "powershell": "powershell",
    }

    if not lang:
        _CLEAN_STATS["detect_code_block_start_fired"] += 1
        return "text"

    _CLEAN_STATS["detect_code_block_start_fired"] += 1
    return aliases.get(lang, lang)

def to_markdown(chapter: dict, image_positions: list = None, next_heading_idx: int = None, book_name: str = "") -> str:
    lines = [f"# {chapter['title']}", ""]

    content = chapter["content"]
    if not content:
        return "\n".join(lines)

    ch_start = chapter.get("heading_doc_index", 0)
    ch_end = next_heading_idx if next_heading_idx is not None else float("inf")

    # =========================
    # Build chapter_images
    # =========================
    chapter_images = []
    if image_positions:
        for item in image_positions:
            para_idx = item[0]

            if para_idx < 0:
                continue

            run_idx = item[1]
            filename = item[3]
            w_px = item[4] if len(item) > 4 else 0
            h_px = item[5] if len(item) > 5 else 0

            if ch_start <= para_idx < ch_end:
                chapter_images.append((para_idx, run_idx, filename, w_px, h_px))

    chapter_images.sort(key=lambda x: (x[0], x[1]))

    assets_path = f"/{book_name}/assets"

    # =========================
    # Build images_by_para_run (correct)
    # =========================
    images_by_para_run = {}

    for p_idx, r_idx, filename, w_px, h_px in chapter_images:
        images_by_para_run.setdefault(p_idx, {}).setdefault(r_idx, []).append(
            (filename, w_px, h_px)
        )

    # =========================
    # Main loop
    # =========================
    i = 0

    while i < len(content):
        item = content[i]
        style = item["style"]
        text = item["text"]
        para_idx = item.get("para_index", -1)
        indent_level = item.get("indent_level", 0)

        lang_marker = _detect_code_block_start(text)

        # =========================
        # Spacing
        # =========================
        if style == "Spacing":
            lines.append("")
            i += 1
            continue

        # =========================
        # Code blocks
        # =========================
        elif lang_marker:
            code_lines = []
            j = i + 1
            closing_found = False

            while j < len(content):
                next_text = content[j]["text"]

                if next_text.strip() == "```":
                    closing_found = True
                    j += 1
                    break

                code_lines.append(next_text)
                j += 1

            if closing_found:
                lines.append(f"```{lang_marker}")
                lines.extend(code_lines)
                lines.append("```")
                lines.append("")

                i = j
                continue
            else:
                lines.append(_preserve_soft_breaks(text))
                lines.append("")
                i += 1
                continue

        # =========================
        # Inline code
        # =========================
        # Rely on the style set by ingest.py, which classifies paragraphs
        # as code based on monospace fonts and language detection. The old
        # fallback of "text starts with backtick" misfired on Markdown-
        # about-Markdown content (books that discuss code syntax).
        elif style == "code":
            lines.append(text)
            lines.append("")
            i += 1
            continue

        # =========================
        # Headings
        # =========================
        elif "Heading 2" in style:
            lines.append(f"## {text}")

        elif "Heading 3" in style:
            lines.append(f"### {text}")

        # =========================
        # Lists
        # =========================
        # ingest.py already injects the correct list prefix (bullet char or
        # number) and indentation into the text, based on the numbering data
        # it extracted from the .docx file. We pass the text through as-is.
        # Previously this branch added another "- " prefix, which produced
        # duplicate markers like "- •" or "- 1)" that _clean_markdown_final
        # then had to clean up with regex (rules 02 and 03).
        elif "List" in style:
            lines.append(text)

        # =========================
        # Tables
        # =========================
        elif style == "Table":
            lines.append(text)

        # =========================
        # Paragraph + images (INLINE)
        # =========================
        else:
            para_images = images_by_para_run.get(para_idx, {})

            # Images BEFORE text (run 0)
            if 0 in para_images:
                for img_filename, w_px, h_px in para_images[0]:
                    if w_px > 0 and h_px > 0:
                        lines.append(
                            f'<img src="{assets_path}/{img_filename}" alt="{img_filename}" width="{w_px}" height="{h_px}" />'
                        )
                    else:
                        lines.append(f"![{img_filename}]({assets_path}/{img_filename})")
                    lines.append("")

            # Text
            lines.append(_preserve_soft_breaks(text))

            # Images AFTER text (run order)
            for r_idx in sorted(para_images.keys()):
                if r_idx == 0:
                    continue

                for img_filename, w_px, h_px in para_images[r_idx]:
                    if w_px > 0 and h_px > 0:
                        lines.append(
                            f'<img src="{assets_path}/{img_filename}" alt="{img_filename}" width="{w_px}" height="{h_px}" />'
                        )
                    else:
                        lines.append(f"![{img_filename}]({assets_path}/{img_filename})")
                    lines.append("")

        lines.append("")
        i += 1

    result = "\n".join(lines)
    result = _clean_markdown_preserving_code(result)

    return result